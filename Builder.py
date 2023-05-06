
from Settings import Settings
from docx import Document
class Builder():

    def __init__(self, insertData, settings: Settings) -> None:
        self.insertData = insertData
        self.settings = settings
        self.PathToDeploy = settings.ini['PATHS']["Deploy"]
        self.PathToSh = settings.ini["PATHS"]["sh"]
        self.out = Document(self.PathToSh)
        self.tags = {}

    def getTags(self):
        fullText = []
        for para in self.out.paragraphs:
            fullText.append(para.text)
        self.tags = {}
        for text in fullText:
            if "{" in text and "}" in text:
                self.tags.update({text.split('{')[1].split('}')[0]: None})

    def checkInputData(self):

        for tag in self.tags.keys():
            if not (tag in self.insertData.keys()):
                raise Exception(f"{tag} not finded")
        for tag in self.insertData.keys():
            if not (tag in self.tags.keys()):
                print(Warning(f"{tag} not used in document"))

    def build(self):

        self.getTags()
        self.checkInputData()
        self.replaceTags()
        self.createTable1(self.insertData["table1"])
        self.createTable2(self.insertData["table2"])
        self.createTable3(self.insertData["table3"])

        self.out.save(self.PathToDeploy + "/Deploy.docx")

    def replaceTags(self):
        for para in self.out.paragraphs:
            text = para.text
            if "{" in text and "}" in text:
                for key in self.insertData.keys():
                    if key in text:
                        para.text = para.text.replace(f"{'{'}{key}{'}'}", self.insertData[key])

    def createRowForTable1(self, table, t):
        lenRows = len(table.rows)
        for n in range(t):
            table.add_row()
        for n in range(3):
            for i in range(t):
                table.cell(row_idx=lenRows + i, col_idx=n).merge(table.cell(row_idx=lenRows + n, col_idx=n))
        return table

    def createTable1(self, table1):
        table = self.out.tables[0]
        ind = 1
        for RPD in table1:
            lenRows = len(table.rows)
            table = self.createRowForTable1(table, len(RPD["indexes"]))
            table.cell(row_idx=lenRows + 2, col_idx=0).text = str(ind)
            table.cell(row_idx=lenRows + 2, col_idx=1).text = str(RPD['codeRPDs'])
            table.cell(row_idx=lenRows + 2, col_idx=2).text = str(RPD['soder_komp'])
            i = 0
            for index in RPD["indexes"]:
                table.cell(row_idx=lenRows + i, col_idx=3).text = str(index["index"])
                table.cell(row_idx=lenRows + i, col_idx=4).text = str(index["rez"])
                table.cell(row_idx=lenRows + i, col_idx=5).text = str(index["name"])
                i += 1
            ind += 1
        table.style = 'Table Grid'

    def createTable2(self, table2):
        table = self.out.tables[1]
        table.style = 'Table Grid'
        for theme in table2:
            table.add_row()
            indexRow = len(table.rows)
            i = 0
            for column in theme.keys():
                table.cell(indexRow - 1, i).text = theme[column]
                i += 1

    def createTable3(self, table3):
        table = self.out.tables[2]
        table.style = 'Table Grid'
        for sred in table3:

            for i in range(len(sred["krit"])):
                table.add_row()
                if i > 0:
                    table.cell(len(table.rows) - 1, 0).merge(table.cell(len(table.rows) - 2, 0))
                table.cell(len(table.rows) - 1, 1).text = str(sred["krit"][i]['body'])
                table.cell(len(table.rows) - 1, 2).text = str(sred["krit"][i]['mark'])

            table.cell(len(table.rows) - 1, 0).text = sred["sred"]
