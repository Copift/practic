from Settings import Settings
from PyPDF2 import PdfReader
from tabula import read_pdf
from docx import Document
from tabulate import tabulate
class AgregatorRpd():
    def __init__(self,settings:Settings) -> None:
        self.settings = settings
        self.listRPds=[]
        self.pathToRPDs=self.settings.RPDs
    def getListRPDs(self):
        for file in self.pathToRPDs:
            print(file)
    def getOS_RPD(self,doc):
        for row in range(len(doc.tables[6].rows)):
            if doc.tables[6].cell(row_idx=row, col_idx=0).text=="Формы текущего контроля":
                return doc.tables[6].cell(row_idx=row, col_idx=1).text

    def getIndicators(self,doc):
        table = doc.tables[3]

        table1 = []

        for i in range(1, len(table.rows)):
            table1.append(table.cell(row_idx=i, col_idx=1).text)

        return table1

    def readTable1(self,doc):
        table=doc.tables[3]
        i=1
        table1=[]
        komp={}
        for i in range(1,len(table.rows)):


            if table.cell(row_idx=i, col_idx=0).text in komp.keys():
                    komp[table.cell(row_idx=i, col_idx=0).text].append({'index':table.cell(row_idx=i, col_idx=1).text,
                                'rez': table.cell(row_idx=i, col_idx=2).text,'name':''
                                })
            else:

                    komp.update({table.cell(row_idx=i, col_idx=0).text:[{'index':table.cell(row_idx=i, col_idx=1).text,
                                'rez': table.cell(row_idx=i, col_idx=2).text, 'name':''
                                }]})

        for  key in komp.keys():
            table1.append({"codeRPDs": key.split(' ')[0],
                           "soder_komp": key.replace(key.split(' ')[0],''),
                           "indexes":komp[key]})

        return table1

    def readTable2(self,doc):


        result = []
        for table in doc.tables:
            if "Раздел дисциплины" in table.cell(row_idx=0, col_idx=1).text:
                for row in range(3, len(table.rows) - 1):
                    result.append({"theme": table.cell(row_idx=row, col_idx=1).text,
                                   "code": "",
                                   "rez": f"",
                                   "mark": "",
                                   "krit": "",
                                   "TK": "",
                                   "PA": ""})
        return result

    def getTable8_2(self,doc):
        result = []
        for table in doc.tables:
            if "Вид контроля" == table.cell(row_idx=0, col_idx=1).text and "Компетенции, компоненты которых контролируются" == table.cell(row_idx=0, col_idx=2).text:
                for row in range(1, len(table.rows) ):
                    result.append(table.cell(row_idx=row, col_idx=1).text)
        return result


    def readDocx(self, file):
        doc=Document(file)



        self.readTable1(doc)

        insertData = {
                      'input.inst': doc.tables[0].cell(row_idx=0,col_idx=1).text.split('\n')[1].replace('Декан факультета ',''),
                      'upload.department': doc.paragraphs[7].text.replace('Кафедра',''),
                      'input.discipline': doc.tables[1].cell(row_idx=0,col_idx=3).text,
                      'input.dirPod': doc.tables[1].cell(row_idx=2,col_idx=2).text,
                      'input.naprav': doc.tables[1].cell(row_idx=4,col_idx=4).text,
                      'input.developers': doc.tables[len(doc.tables)-1].cell(row_idx=0, col_idx=4).text,
                      'upload.OS_RPD': self.getOS_RPD(doc),
                      'table8_2':self.getTable8_2(doc),
                      'table1':self.readTable1(doc),
                       'table2': self.readTable2(doc),

        }
        insertTable1_3=[]
        for i in range(7):
            krit = []
            for n in range(7 - i):
                krit.append({
                    "body": f"body_test_{n + 1}",
                    "mark": n
                })
            insertTable1_3.append({
                "sred": f"sred_{i + 1}",
                "krit": krit,
            })
        insertData.update({'table3': insertTable1_3})
        print(insertData)
        return insertData








